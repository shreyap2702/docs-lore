import gc
from fastapi import FastAPI, Depends, HTTPException, status
import psutil
from pydantic import BaseModel
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.staticfiles import StaticFiles
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pinecone import Pinecone, ServerlessSpec
from langchain_pinecone import PineconeVectorStore
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnableLambda
from langchain.chains import create_retrieval_chain
from pypdf import PdfReader
import mimetypes
import asyncio
import requests
import logging
import io
import os
import docx
import email
from email import policy
from dotenv import load_dotenv
import hashlib
import pickle
import json
from sentence_transformers import SentenceTransformer
import tempfile
from pdfminer.high_level import extract_text as pdfminer_extract_text

load_dotenv()

# --- Render-specific configurations ---
IS_RENDER = os.getenv("RENDER", "false").lower() == "true"
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB for Render
REQUEST_TIMEOUT = 120  # 2 minutes for large files
CHUNK_SIZE = 8192  # 8KB chunks for streaming

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_ENVIRONMENT = os.getenv("PINECONE_ENVIRONMENT")
PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX_NAME", "hackathon-policy-docs")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- Memory management for Render ---
if IS_RENDER:
    logger.info("Running on Render - enabling memory optimizations")
    # Reduce chunk size for memory efficiency
    CHUNK_SIZE = 4096  # 4KB chunks on Render
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB limit on Render

# --- Enhanced Caching System ---
processed_documents_cache = set()
embedding_cache_dir = "embedding_cache"
os.makedirs(embedding_cache_dir, exist_ok=True)

# --- Local Embeddings for API Quota Reduction ---
USE_LOCAL_EMBEDDINGS = False  # Set to False to use Google embeddings for now
if USE_LOCAL_EMBEDDINGS:
    local_embeddings_model = SentenceTransformer('all-mpnet-base-v2')  # 768 dimensions, matches your index
    print("Using local embeddings (all-mpnet-base-v2) to reduce API quotas")
else:
    local_embeddings_model = None

# --- Global Initializations ---
# Enhanced chunking with semantic awareness
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=800,  # Slightly smaller for better precision
    chunk_overlap=150,  # Reduced overlap
    length_function=len,
    separators=["\n\n", "\n", ". ", " ", ""]
)

llm = ChatGoogleGenerativeAI(model="models/gemini-2.5-flash-lite", temperature=0.2, google_api_key=GOOGLE_API_KEY)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)

pc = Pinecone(api_key=PINECONE_API_KEY, environment=PINECONE_ENVIRONMENT)

if PINECONE_INDEX_NAME not in pc.list_indexes().names():
    pc.create_index(
        name=PINECONE_INDEX_NAME,
        dimension=768, # Dimension for "models/embedding-001" is typically 768
        metric="cosine",
        spec=ServerlessSpec(cloud='aws', region='us-west-2')
    )
    logger.info(f"Created new Pinecone index: {PINECONE_INDEX_NAME}")
else:
    logger.info(f"Pinecone index '{PINECONE_INDEX_NAME}' already exists.")

# --- Enhanced Embedding Cache Class ---
class EmbeddingCache:
    def __init__(self, cache_dir="embedding_cache"):
        self.cache_dir = cache_dir
        os.makedirs(cache_dir, exist_ok=True)
    
    def get_embedding(self, text, use_local=True):
        text_hash = hashlib.md5(text.encode()).hexdigest()
        cache_file = os.path.join(self.cache_dir, f"{text_hash}.pkl")
        
        if os.path.exists(cache_file):
            with open(cache_file, 'rb') as f:
                return pickle.load(f)
        
        # Generate embedding
        if use_local and local_embeddings_model:
            embedding = local_embeddings_model.encode(text).tolist()
        else:
            embedding = embeddings.embed_query(text)
        
        # Cache it
        with open(cache_file, 'wb') as f:
            pickle.dump(embedding, f)
        
        return embedding

embedding_cache = EmbeddingCache()

# --- Custom Embedding Class for Local Models ---
class LocalEmbeddings:
    def __init__(self, model_name='all-mpnet-base-v2'):
        self.model = SentenceTransformer(model_name)
    
    def embed_query(self, text):
        return self.model.encode(text).tolist()
    
    def embed_documents(self, texts):
        return self.model.encode(texts).tolist()
    
    async def aembed_query(self, text):
        return self.model.encode(text).tolist()
    
    async def aembed_documents(self, texts):
        return self.model.encode(texts).tolist()

# --- Enhanced Vector Store with Local Embeddings ---
if USE_LOCAL_EMBEDDINGS:
    local_embeddings = LocalEmbeddings()
    vectorstore = PineconeVectorStore(
        index=pc.Index(PINECONE_INDEX_NAME), 
        embedding=local_embeddings,
        namespace="hackrx-2025"
    )
else:
    vectorstore = PineconeVectorStore(index=pc.Index(PINECONE_INDEX_NAME), embedding=embeddings, namespace="hackrx-2025")

# --- Shorter, More Focused RAG Prompt (25% shorter responses) ---
rag_prompt_template = ChatPromptTemplate.from_template("""
You are a document analyzer. Answer the user's question based on the provided context.

Guidelines:
- Be concise and direct
- Focus on the most relevant information
- Use clean, professional formatting
- Avoid newline characters, bullet points, asterisks, or special formatting
- If information is missing, state "Information not available in the document"
- For insurance/legal documents, provide clear, structured answers
- Use simple text format without markdown or special characters

Context: {context}
Question: {input}

Provide a clear, professional response:
""")

document_chain = create_stuff_documents_chain(llm, rag_prompt_template)

# --- Enhanced Retrieval with Better Parameters ---
retriever = vectorstore.as_retriever(
    search_kwargs={
        "k": 6,  # Reduced for more focused results
        "score_threshold": 0.35,  # Slightly lower threshold for better coverage
        "include_metadata": True  # Include metadata for better context
    }
)

# --- Multi-stage retrieval for better results ---
def enhanced_retrieval(question: str):
    # Stage 1: Get initial results
    docs = retriever.invoke(question)  # Fixed: Using .invoke() instead of deprecated .get_relevant_documents()
    
    # Stage 2: Filter and re-rank based on relevance
    if docs:
        # Simple re-ranking based on keyword overlap
        scored_docs = []
        question_words = set(question.lower().split())
        
        for doc in docs:
            doc_words = set(doc.page_content.lower().split())
            overlap = len(question_words.intersection(doc_words))
            score = overlap / max(len(question_words), 1)
            scored_docs.append((doc, score))
        
        # Sort by score and take top results
        scored_docs.sort(key=lambda x: x[1], reverse=True)
        final_docs = [doc for doc, score in scored_docs[:4]]  # Top 4 most relevant
        
        return final_docs
    
    return docs

rag_chain = create_retrieval_chain(retriever, document_chain)

# --- End Global Initializations ---

app = FastAPI(title="LLM Query-Retrieval System")

if os.path.isdir("static"):
    app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/health")
async def health_check():
    """Health check endpoint for Render"""
    return {
        "status": "healthy",
        "timestamp": "2025-08-01T21:30:00Z",
        "service": "DocLore API",
        "version": "1.0.0"
    }

EXPECTED_AUTH_TOKEN = "10fbd8807c6d9b5a37028c3eb4bd885959f06d006aedd2dc5ba64c5ccea913c0"
security = HTTPBearer()


class QueryRequest(BaseModel):
    documents: str
    questions: list[str]


class QueryResponse(BaseModel):
    answers: list[str]


async def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    if credentials.scheme != "Bearer" or credentials.credentials != EXPECTED_AUTH_TOKEN:
        logger.warning("Authentication failed: Invalid token.")
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid or missing authentication token"
        )
    return True


def validate_text_input(text_input):
    """Validate and clean text input for embeddings"""
    if text_input is None:
        return ""
    if not isinstance(text_input, str):
        text_input = str(text_input)
    return text_input.strip()

def log_memory_usage(operation):
    """Log memory usage for monitoring"""
    try:
        memory_mb = psutil.Process().memory_info().rss / 1024 / 1024
        logger.info(f"Memory after {operation}: {memory_mb:.2f} MB")
    except Exception as e:
        logger.warning(f"Could not log memory usage: {e}")

def check_memory_limit():
    """Memory circuit breaker - prevents crashes on Render"""
    try:
        memory_mb = psutil.Process().memory_info().rss / 1024 / 1024
        if memory_mb > 450:  # 88% of 512MB limit
            logger.error(f"CRITICAL MEMORY: {memory_mb}MB - Aborting request")
            gc.collect()
            raise HTTPException(status_code=503, detail="Service temporarily overloaded")
        return memory_mb
    except Exception as e:
        logger.warning(f"Could not check memory: {e}")
        return 0

def emergency_cleanup():
    """Aggressive memory cleanup when approaching limits"""
    try:
        # Force garbage collection multiple times
        for i in range(3):
            collected = gc.collect()
            logger.info(f"Emergency cleanup round {i+1}: collected {collected} objects")
    except Exception as e:
        logger.warning(f"Emergency cleanup failed: {e}")

def safe_memory_check(operation="operation"):
    """Safe memory check with automatic cleanup"""
    try:
        memory_mb = psutil.Process().memory_info().rss / 1024 / 1024
        logger.info(f"Memory after {operation}: {memory_mb:.2f} MB")
        
        if memory_mb > 450:
            logger.error(f"CRITICAL: {memory_mb}MB - Aborting")
            emergency_cleanup()
            raise HTTPException(status_code=503, detail="Memory limit exceeded")
        
        if memory_mb > 400:
            logger.warning(f"HIGH MEMORY: {memory_mb}MB - Forcing cleanup")
            emergency_cleanup()
            
        return memory_mb
    except Exception as e:
        logger.warning(f"Memory check failed: {e}")
        return 0

def get_dynamic_batch_size():
    """Adjust batch size based on current memory"""
    try:
        memory_mb = psutil.Process().memory_info().rss / 1024 / 1024
        
        if memory_mb > 400:
            return 2  # Very small batches
        elif memory_mb > 350:
            return 3  # Small batches
        else:
            return 5  # Normal batches
    except Exception as e:
        logger.warning(f"Could not get dynamic batch size: {e}")
        return 3  # Default to small batches

def validate_document_size(file_bytes, doc_url):
    """Check if document can be processed safely"""
    try:
        file_size_mb = len(file_bytes) / (1024 * 1024)
        current_memory = psutil.Process().memory_info().rss / 1024 / 1024
        
        # Conservative estimate: document expands 6-10x in memory
        estimated_memory = file_size_mb * 10
        available_memory = 512 - current_memory  # Render limit
        
        if estimated_memory > available_memory:
            logger.warning(f"Large document detected: {file_size_mb:.1f}MB - will use chunked processing")
            return False  # Signal to use chunked processing
            
        logger.info(f"Document size: {file_size_mb:.1f}MB, Estimated memory: {estimated_memory:.1f}MB, Available: {available_memory:.1f}MB")
        return True  # Safe to process normally
        
    except Exception as e:
        logger.warning(f"Document size validation failed: {e}")
        return True  # Default to normal processing

async def process_large_document_chunked(file_bytes: bytes, file_type: str) -> str:
    """Process large documents in chunks to avoid memory limits"""
    try:
        logger.info(f"Starting chunked processing for {file_type} document")
        
        if file_type == "pdf":
            return await parse_pdf_chunked(file_bytes)
        elif file_type == "docx":
            return await parse_docx_chunked(file_bytes)
        elif file_type == "eml":
            return await parse_email_chunked(file_bytes)
        else:
            raise HTTPException(status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE, detail=f"Unsupported document type: {file_type}")
            
    except Exception as e:
        logger.error(f"Chunked processing failed: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to process large document: {str(e)}")

async def parse_pdf_chunked(data: bytes) -> str:
    """Parse PDF in chunks to avoid memory limits"""
    logger.info("Starting chunked PDF parsing")
    
    try:
        reader = await asyncio.to_thread(PdfReader, io.BytesIO(data))
        text_parts = []
        
        # Process pages in smaller batches
        page_batch_size = 10  # Process 10 pages at a time
        total_pages = len(reader.pages)
        
        for batch_start in range(0, total_pages, page_batch_size):
            batch_end = min(batch_start + page_batch_size, total_pages)
            batch_pages = reader.pages[batch_start:batch_end]
            
            logger.info(f"Processing PDF pages {batch_start+1}-{batch_end} of {total_pages}")
            
            for page in batch_pages:
                page_text = await asyncio.to_thread(page.extract_text)
                if page_text:
                    text_parts.append(page_text.strip())
            
            # Memory cleanup after each batch
            del batch_pages
            gc.collect()
            safe_memory_check(f"PDF batch {batch_start//page_batch_size + 1}")
        
        text = "\n".join(text_parts)
        text = validate_text_input(text)
        
        # Final cleanup
        del reader, text_parts
        gc.collect()
        safe_memory_check("PDF chunked parsing complete")
        
        logger.info(f"Successfully parsed PDF in chunks. Extracted {len(text)} characters.")
        return text
        
    except Exception as e:
        logger.error(f"Chunked PDF parsing error: {e}", exc_info=True)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"PDF parsing error: {e}")

async def parse_docx_chunked(data: bytes) -> str:
    """Parse DOCX in chunks to avoid memory limits"""
    logger.info("Starting chunked DOCX parsing")
    
    try:
        doc = await asyncio.to_thread(docx.Document, io.BytesIO(data))
        text_parts = []
        
        # Process paragraphs in batches
        para_batch_size = 50  # Process 50 paragraphs at a time
        total_paragraphs = len(doc.paragraphs)
        
        for batch_start in range(0, total_paragraphs, para_batch_size):
            batch_end = min(batch_start + para_batch_size, total_paragraphs)
            batch_paragraphs = doc.paragraphs[batch_start:batch_end]
            
            logger.info(f"Processing DOCX paragraphs {batch_start+1}-{batch_end} of {total_paragraphs}")
            
            for para in batch_paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Memory cleanup after each batch
            del batch_paragraphs
            gc.collect()
            safe_memory_check(f"DOCX batch {batch_start//para_batch_size + 1}")
        
        text = "\n".join(text_parts)
        text = validate_text_input(text)
        
        # Final cleanup
        del doc, text_parts
        gc.collect()
        safe_memory_check("DOCX chunked parsing complete")
        
        logger.info(f"Successfully parsed DOCX in chunks. Extracted {len(text)} characters.")
        return text
        
    except Exception as e:
        logger.error(f"Chunked DOCX parsing error: {e}", exc_info=True)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"DOCX parsing error: {e}")

async def parse_email_chunked(data: bytes) -> str:
    """Parse email in chunks to avoid memory limits"""
    logger.info("Starting chunked email parsing")
    
    try:
        msg = await asyncio.to_thread(email.message_from_bytes, data, policy=policy.default)
        main_body = msg.get_body()
        text_parts = []
        
        if main_body:
            # Process email parts in smaller chunks
            parts = list(main_body.iter_attachments())
            part_batch_size = 5  # Process 5 parts at a time
            
            for batch_start in range(0, len(parts), part_batch_size):
                batch_end = min(batch_start + part_batch_size, len(parts))
                batch_parts = parts[batch_start:batch_end]
                
                logger.info(f"Processing email parts {batch_start+1}-{batch_end} of {len(parts)}")
                
                for part in batch_parts:
                    if part.get_content_type() == 'text/plain':
                        part_text = part.get_content()
                        if part_text:
                            text_parts.append(part_text)
                
                # Memory cleanup after each batch
                del batch_parts
                gc.collect()
                safe_memory_check(f"Email batch {batch_start//part_batch_size + 1}")
            
            # Handle main body content
            if main_body.get_content_type() == 'text/plain':
                main_text = main_body.get_content()
                if main_text:
                    text_parts.append(main_text)
        
        text = "\n".join(text_parts)
        text = validate_text_input(text)
        
        # Final cleanup
        del msg, main_body, text_parts
        gc.collect()
        safe_memory_check("Email chunked parsing complete")
        
        logger.info(f"Successfully parsed email in chunks. Extracted {len(text)} characters.")
        return text
        
    except Exception as e:
        logger.error(f"Chunked email parsing error: {e}", exc_info=True)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Email parsing error: {e}")


def stream_pdf_to_tempfile(pdf_url, chunk_size=1024*1024):
    """Stream a PDF from a URL to a temporary file on disk."""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            with requests.get(pdf_url, stream=True, timeout=REQUEST_TIMEOUT) as r:
                r.raise_for_status()
                for chunk in r.iter_content(chunk_size=chunk_size):
                    if chunk:
                        tmp_file.write(chunk)
            return tmp_file.name
    except Exception as e:
        logger.error(f"Failed to stream PDF to tempfile: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to stream PDF: {e}")

def extract_text_from_pdf_file(pdf_path):
    """Extract text from a PDF file using pdfminer.six."""
    try:
        text = pdfminer_extract_text(pdf_path)
        return validate_text_input(text)
    except Exception as e:
        logger.error(f"Failed to extract text from PDF file: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to extract text from PDF: {e}")
    finally:
        try:
            os.remove(pdf_path)
        except Exception as cleanup_err:
            logger.warning(f"Failed to remove temp PDF file: {cleanup_err}")


def process_document_in_chunks(text: str, source_name: str, chunk_size_chars: int = 50000):
    """Process large documents in smaller chunks to avoid memory issues"""
    try:
        logger.info(f"Processing large document in chunks: {len(text)} characters")
        
        # Split text into manageable chunks
        text_chunks = []
        for i in range(0, len(text), chunk_size_chars):
            chunk_text = text[i:i + chunk_size_chars]
            text_chunks.append(chunk_text)
        
        logger.info(f"Split document into {len(text_chunks)} chunks")
        
        all_chunks = []
        for i, chunk_text in enumerate(text_chunks):
            # Process each chunk separately
            chunks = chunk_text(chunk_text, f"{source_name}_chunk_{i}")
            all_chunks.extend(chunks)
            
            # Memory cleanup after each chunk
            del chunks
            gc.collect()
            safe_memory_check(f"chunk {i+1} processing")
            
            # Stop if memory is getting too high
            if check_memory_limit() > 400:
                logger.warning(f"Stopping chunk processing at chunk {i+1} due to memory constraints")
                break
        
        logger.info(f"Successfully processed {len(all_chunks)} total chunks")
        return all_chunks
        
    except Exception as e:
        logger.error(f"Chunked processing failed: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to process large document"
        )


async def download_file_content(url: str) -> tuple[bytes, str]:
    logger.info(f"Downloading from {url}")
    try:
        # Enhanced download with better error handling and size limits
        response = await asyncio.to_thread(
            requests.get, 
            url, 
            stream=True, 
            timeout=REQUEST_TIMEOUT,  # Use Render-specific timeout
            headers={'User-Agent': 'Mozilla/5.0 (compatible; DocLore/1.0)'}
        )
        response.raise_for_status()
        
        # Check file size before downloading
        content_length = response.headers.get('content-length')
        if content_length:
            file_size = int(content_length)
            if file_size > MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                    detail=f"File too large: {file_size / (1024*1024):.1f}MB. Maximum allowed: {MAX_FILE_SIZE / (1024*1024)}MB"
                )
        
        # Stream download for large files with memory monitoring
        content = b""
        chunk_count = 0
        
        for chunk in response.iter_content(chunk_size=CHUNK_SIZE):
            if chunk:  # Filter out keep-alive chunks
                content += chunk
                chunk_count += 1
                
                # Check memory every 10 chunks to avoid excessive checking
                if chunk_count % 10 == 0:
                    current_memory = psutil.Process().memory_info().rss / 1024 / 1024
                    if current_memory > 450:  # Critical memory limit
                        logger.error(f"Memory limit reached during download: {current_memory:.1f}MB")
                        raise HTTPException(
                            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                            detail="Service temporarily overloaded during download"
                        )
                
                # Check if we're exceeding file size limits
                if len(content) > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                        detail=f"File download exceeded size limit of {MAX_FILE_SIZE / (1024*1024)}MB"
                    )
        
        content_type = response.headers.get("Content-Type", "")

        # Validate downloaded content
        if not isinstance(content, bytes) or len(content) == 0:
            logger.error(f"Downloaded content for {url} is not bytes or is empty. Type: {type(content)}, Length: {len(content) if content else 0}")
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Downloaded file from {url} is empty or malformed."
            )

        logger.info(f"Successfully downloaded {len(content)} bytes from {url} in {chunk_count} chunks")
        return content, content_type
        
    except requests.exceptions.Timeout:
        logger.error(f"Download timed out for {url}")
        raise HTTPException(status_code=status.HTTP_408_REQUEST_TIMEOUT, detail="Download timed out")
    except requests.exceptions.RequestException as e:
        logger.error(f"Request error for {url}: {e}")
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"Request error: {e}")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error downloading {url}: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Download failed: {str(e)}")


def detect_file_type(url: str, content_type: str) -> str:
    ext_type, _ = mimetypes.guess_type(url)
    mime = content_type.lower()
    if "pdf" in mime or ext_type == "application/pdf":
        return "pdf"
    elif "word" in mime or ext_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return "docx"
    elif "rfc822" in mime or ext_type == "message/rfc822":
        return "eml"
    return "unknown"


async def parse_pdf(data: bytes) -> str:
    logger.info("Attempting to parse PDF bytes into text.")
    try:
        reader = await asyncio.to_thread(PdfReader, io.BytesIO(data))
        text_parts = []
        for page in reader.pages:
            page_text = await asyncio.to_thread(page.extract_text)
            if page_text:
                text_parts.append(page_text.strip())
        
        text = "\n".join(text_parts)
        # Validate extracted text
        text = validate_text_input(text)
        logger.info(f"Successfully parsed PDF. Extracted {len(text)} characters.")
        return text
    except Exception as e:
        logger.error(f"PDF parsing error: {e}", exc_info=True)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"PDF parsing error: {e}")


async def parse_docx(data: bytes) -> str:
    logger.info("Attempting to parse DOCX bytes into text.")
    try:
        doc = await asyncio.to_thread(docx.Document, io.BytesIO(data))
        text_parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(text_parts)
        # Validate extracted text
        text = validate_text_input(text)
        logger.info(f"Successfully parsed DOCX. Extracted {len(text)} characters.")
        return text
    except Exception as e:
        logger.error(f"DOCX parsing error: {e}", exc_info=True)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"DOCX parsing error: {e}")


async def parse_email(data: bytes) -> str:
    logger.info("Attempting to parse email bytes into text.")
    try:
        msg = await asyncio.to_thread(email.message_from_bytes, data, policy=policy.default)
        main_body = msg.get_body()
        text = ""
        if main_body:
            for part in main_body.iter_attachments():
                if part.get_content_type() == 'text/plain':
                    text = part.get_content()
                    break
            if not text and main_body.get_content_type() == 'text/plain':
                text = main_body.get_content()
            elif not text and main_body.get_content_type() == 'text/html':
                text = main_body.get_content()
                logger.warning("Extracted HTML content from email. Consider using HTML parser for cleaner text.")
        
        # Validate extracted text
        text = validate_text_input(text)
        logger.info(f"Successfully parsed email. Extracted {len(text)} characters.")
        return text
    except Exception as e:
        logger.error(f"Email parsing error: {e}", exc_info=True)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Email parsing error: {e}")


def chunk_text(text: str, source_name: str) -> list[Document]:
    # Validate input text
    text = validate_text_input(text)
    if not text:
        logger.warning(f"Empty text provided for chunking from source: {source_name}")
        return []
    
    # Enhanced chunking with better structure preservation
    chunks = text_splitter.create_documents([text])
    valid_chunks = []
    
    for i, chunk in enumerate(chunks):
        # Ensure chunk content is valid
        if hasattr(chunk, 'page_content'):
            chunk.page_content = validate_text_input(chunk.page_content)
            if chunk.page_content:  # Only add non-empty chunks
                # Enhanced metadata for better retrieval
                chunk.metadata.update({
                    "source": source_name,
                    "chunk_index": i,
                    "chunk_type": "semantic",
                    "length": len(chunk.page_content),
                    "word_count": len(chunk.page_content.split())
                })
                valid_chunks.append(chunk)
        else:
            logger.warning(f"Invalid chunk format: {type(chunk)}")
    
    logger.info(f"Created {len(valid_chunks)} valid chunks from {len(chunks)} total chunks")
    return valid_chunks


@app.post("/api/v1/hackrx/run", response_model=QueryResponse)
async def run_submission(request: QueryRequest, auth: bool = Depends(verify_token)):
    logger.info(f"Received request for document: {request.documents}")
    logger.info(f"Questions: {request.questions}")

    # Initial memory check - fail fast if already high
    initial_memory = check_memory_limit()
    logger.info(f"Initial memory: {initial_memory:.2f}MB")

    try:
        if request.documents not in processed_documents_cache:
            logger.info(f"Document {request.documents} not in cache. Proceeding with ingestion.")

            file_type = detect_file_type(request.documents, "application/pdf")  # Only for PDF streaming
            if file_type == "pdf":
                # Use streaming for all PDFs (or add logic for large PDFs only)
                pdf_path = await asyncio.to_thread(stream_pdf_to_tempfile, request.documents)
                text = await asyncio.to_thread(extract_text_from_pdf_file, pdf_path)
                if not text.strip():
                    logger.error(f"No text extracted from document: {request.documents}")
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                        detail="No text content could be extracted from the document"
                    )
                chunks_for_pinecone = chunk_text(text, request.documents)
                if not chunks_for_pinecone:
                    logger.error(f"No valid chunks created from document: {request.documents}")
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                        detail="No valid text chunks could be created from the document"
                    )
                batch_size = get_dynamic_batch_size()
                max_retries = 3
                logger.info(f"Using dynamic batch size: {batch_size} (memory-based)")
                for i in range(0, len(chunks_for_pinecone), batch_size):
                    batch = chunks_for_pinecone[i:i + batch_size]
                    for attempt in range(max_retries):
                        try:
                            await asyncio.to_thread(vectorstore.add_documents, batch)
                            logger.info(f"Processed batch {i//batch_size + 1}/{(len(chunks_for_pinecone) + batch_size - 1)//batch_size}")
                            del batch
                            gc.collect()
                            safe_memory_check(f"batch {i//batch_size + 1}")
                            break
                        except Exception as e:
                            if "Max retries exceeded" in str(e) or "Failed to resolve" in str(e):
                                if attempt < max_retries - 1:
                                    logger.warning(f"Network error on attempt {attempt + 1}, retrying in 2 seconds...")
                                    await asyncio.sleep(2)
                                else:
                                    logger.error(f"Failed to upload batch after {max_retries} attempts: {e}")
                                    raise HTTPException(
                                        status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                                        detail="Pinecone service temporarily unavailable. Please try again."
                                    )
                            else:
                                raise e
                del chunks_for_pinecone
                gc.collect()
                safe_memory_check("document processing complete")
                logger.info(f"Upserted chunks to Pinecone for {request.documents}.")
                processed_documents_cache.add(request.documents)
            else:
                # ... existing code for non-PDF documents ...
                file_bytes, content_type = await download_file_content(request.documents)
                can_process_normally = validate_document_size(file_bytes, request.documents)
                file_type = detect_file_type(request.documents, content_type)
                logger.info(f"DEBUG: Before parsing - file_bytes type: {type(file_bytes)}")
                logger.info(f"DEBUG: Before parsing - file_bytes length: {len(file_bytes)} bytes")
                logger.info(f"DEBUG: Before parsing - Detected file_type: {file_type}")
                text = ""
                if can_process_normally:
                    if file_type == "docx":
                        text = await parse_docx(file_bytes)
                    elif file_type == "eml":
                        text = await parse_email(file_bytes)
                    else:
                        logger.warning(f"Unsupported document type detected: {file_type} for URL {request.documents}")
                        raise HTTPException(status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE, detail=f"Unsupported document type: {file_type}")
                    del file_bytes
                    gc.collect()
                    safe_memory_check("document parsing")
                else:
                    logger.info("Using chunked processing for large document")
                    text = await process_large_document_chunked(file_bytes, file_type)
                if not text.strip():
                    logger.error(f"No text extracted from document: {request.documents}")
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                        detail="No text content could be extracted from the document"
                    )
                chunks_for_pinecone = chunk_text(text, request.documents)
                if not chunks_for_pinecone:
                    logger.error(f"No valid chunks created from document: {request.documents}")
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                        detail="No valid text chunks could be created from the document"
                    )
                batch_size = get_dynamic_batch_size()
                max_retries = 3
                logger.info(f"Using dynamic batch size: {batch_size} (memory-based)")
                for i in range(0, len(chunks_for_pinecone), batch_size):
                    batch = chunks_for_pinecone[i:i + batch_size]
                    for attempt in range(max_retries):
                        try:
                            await asyncio.to_thread(vectorstore.add_documents, batch)
                            logger.info(f"Processed batch {i//batch_size + 1}/{(len(chunks_for_pinecone) + batch_size - 1)//batch_size}")
                            del batch
                            gc.collect()
                            safe_memory_check(f"batch {i//batch_size + 1}")
                            break
                        except Exception as e:
                            if "Max retries exceeded" in str(e) or "Failed to resolve" in str(e):
                                if attempt < max_retries - 1:
                                    logger.warning(f"Network error on attempt {attempt + 1}, retrying in 2 seconds...")
                                    await asyncio.sleep(2)
                                else:
                                    logger.error(f"Failed to upload batch after {max_retries} attempts: {e}")
                                    raise HTTPException(
                                        status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                                        detail="Pinecone service temporarily unavailable. Please try again."
                                    )
                            else:
                                raise e
                del chunks_for_pinecone
                gc.collect()
                safe_memory_check("document processing complete")
                logger.info(f"Upserted chunks to Pinecone for {request.documents}.")
                processed_documents_cache.add(request.documents)
        else:
            logger.info(f"Document {request.documents} found in cache. Skipping ingestion steps.")

        answers = []
        for i, question in enumerate(request.questions):
            logger.info(f"Answering question {i+1}/{len(request.questions)}: {question}")

            # Check memory before each question - stop if too high
            if i > 0:  # Skip check for first question
                current_memory = check_memory_limit()
                if current_memory > 400 and i > 2:  # If high memory and processed some questions
                    logger.warning(f"Stopping at question {i+1} due to memory constraints ({current_memory:.1f}MB)")
                    answers.append("Processing stopped due to memory constraints. Please try with fewer questions.")
                    break

            # Validate question input
            clean_question = validate_text_input(question)
            if not clean_question:
                logger.warning(f"Empty or invalid question: {question}")
                answers.append("Invalid or empty question provided.")
                continue

            try:
                # Use enhanced retrieval for better results
                docs = await asyncio.to_thread(enhanced_retrieval, clean_question)
                logger.info(f"Retrieved {len(docs)} documents for question: {clean_question}")
                if docs:
                    logger.info(f"First doc preview: {docs[0].page_content[:200]}...")
                
                result = await rag_chain.ainvoke({"input": clean_question})  # Use 'input' key
                answer = result.get("answer", "No answer found")  # Extract the 'answer' key from the result dict
                answers.append(answer)
                logger.info(f"Generated answer for '{question}': {answer[:100]}...")
                
                # Aggressive memory cleanup after each question
                del docs, result, answer
                gc.collect()
                safe_memory_check(f"question {i+1} processing")
                
            except Exception as e:
                logger.error(f"Error processing question '{question}': {e}")
                answers.append(f"Error processing question: {str(e)}")


        return QueryResponse(answers=answers)

    except HTTPException as e:
        raise e
    except Exception as e:
        logger.exception("An unhandled error occurred during run_submission processing.")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"An internal server error occurred: {str(e)}"
        )